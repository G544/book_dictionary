from docx.enum.section import WD_ORIENTATION
from backend.fb2_parser import Pars_and_tokenize
from docx import Document
from docx.oxml.ns import qn

class Book_Dictionary():

    NUMBER_OF_COLS = '4'
    SPACE_BETWEEN_COLS = '20'
    def __init__(self, file, lang):
        ''' Modes = 'Chapter' - extract only define chapter,
        "chapters" - all but each new chapter conatins only new words,
        "chapters_ex" - all book but divided for chapters
        'Book' - dictionary for all book
        '''
        self.lang = lang
        self.filename = file
        self.book = Pars_and_tokenize(self.filename)
        # Creating the document and setting the main feature - 4 columns and
        self.document = Document()
        section = self.document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), self.NUMBER_OF_COLS)
        cols.set(qn('w:space'), self.SPACE_BETWEEN_COLS)  # Set space between columns to 10 points ->0.01"
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

    def convert(self,end_file, mode="chapter", chapter=1):
        self.end_file = end_file
        # define the number of chapter for mode which convert only one chapter
        if mode == "chapter" and chapter:
            self.chapter = chapter
        if mode == 'chapter':
            self.chapter_mode()
        elif mode == 'chapters':
            self.chapters_mode()
        elif mode == 'chapters_ex':
            self.chapters_mode_ex()
        elif mode == 'book':
            self.book_mode()
        else:
            print('Wrong value for mode')

    # process mode which convert only one chapter
    def chapter_mode(self):
        self._add_text_to_file(self.book.translator(self.chapter, self.lang))
        self._write_to_file()

    # process mode which convert all chapter with all words in chapter
    def chapters_mode_ex(self):
        for i in range(self.book.num_chapters):
            self.document.add_paragraph(str(i) + '\n')
            self._add_text_to_file(self.book.translator(i, self.lang))
        self._write_to_file()

    # process mode which convert all chapter with only new words in chapter
    def chapters_mode(self):
        full_dict = []
        for i in range(self.book.num_chapters):
            full_chapter = self.book.translator(i, self.lang)
            new_chapter = full_chapter # may be we don't need here two variables?
            # Remove words which occured in previous chapters
            list(map(new_chapter.__delitem__, filter(new_chapter.__contains__, full_dict)))
            # Because head letters in each chapter, we added it again
            head_letters = set([s[0].title() for s in list(new_chapter.keys())])
            for letter in head_letters:
                new_chapter[letter] = '  '
            new_chapter = dict(sorted(new_chapter.items(), key=lambda x: x[0].lower()))
            # add new words in dict with all words
            full_dict = list(set(full_dict + list(full_chapter.keys())))
            self.document.add_paragraph(str(i) + '\n')
            self._add_text_to_file(new_chapter)
        self._write_to_file()

    # process mode which convert the whoile book
    def book_mode(self):
        full_dict = {}
        for i in range(self.book.num_chapters):
            full_chapter = self.book.translator(i, self.lang)
            # merged old dict and new chapter's dict
            full_dict = {**full_dict, **full_chapter}
        full_dict = dict(sorted(full_dict.items(), key=lambda x: x[0].lower()))
        self._add_text_to_file(full_dict)
        self._write_to_file()

    #  add new part of the dictionary
    def _add_text_to_file(self, dictionary):
        text = '\n'.join(list(' : '.join(pair) for pair in (dictionary.items())))
        paragraph = self.document.add_paragraph(text)
        paragraph.paragraph_format.space_after = 0

    # save ready dictionary
    def _write_to_file(self):
        self.document.save(self.end_file)
