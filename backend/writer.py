from docx.enum.section import WD_ORIENTATION

from docx import Document
from docx.oxml.ns import qn

class SaveToDoc():

    NUMBER_OF_COLS = '4'
    SPACE_BETWEEN_COLS = '20'
    def __init__(self, filename:str):

        self.filename = filename
        # Creating the document and setting the main feature - 4 columns and
        self.document = Document()
        section = self.document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), self.NUMBER_OF_COLS)
        cols.set(qn('w:space'), self.SPACE_BETWEEN_COLS)  # Set space between columns to 10 points ->0.01"
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

    #  add new part of the dictionary
    def add_text_to_file(self, dictionary:dict):
        text = '\n'.join(list(' : '.join(pair) for pair in (dictionary.items())))
        paragraph = self.document.add_paragraph(text)
        paragraph.paragraph_format.space_after = 0

    def add_chapter(self, chapter_name:str):
        self.document.add_paragraph(chapter_name + '\n')
        
    # save ready dictionary
    def write_to_file(self):
        self.document.save(self.filename)
